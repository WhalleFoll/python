import docx
from docx import Document
import time
from docx.shared import Cm,Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

document = Document()
head = document.add_heading('第一章 绪论', 0)
head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph("    物联网技术的研究背景和意义，该部分主要介绍智能家居和行为分析技术的相关研究背景、意义，概述了本文运用到的关键技术和创新点，交代了论文的内容和结构安排。") #写入若干段落
document.add_heading('1.1 物联网技术的研究背景', level=2)
document.add_paragraph('    物联网技术的发展介绍，该部分主要介绍物联网智能家居系统的起源发展和行为分析技术的研究进展，探讨将行为分析技术运用到智能家居系统中的意义。')
document.add_heading('1.2 物联网技术的研究意义', level=2)
document.add_paragraph('    随着互联网‘5G’人工智能的发展，万物互联时代随之到来。什么是物联网')
img = document.add_picture('tu.png', width=Inches(5.0))
img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
table = document.add_table(rows=9, cols=2)
table.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
for r in range(9):#循环将每一行，每一列都设置为居中
    for c in range(2):
        table.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        if c==1:
            table.cell(r, c).text = '2.2s'
            table.cell(r, c).height = Cm(1)
        else:
            table.cell(r, c).text = '空调打开时间'
            table.cell(r, c).height = Cm(1)
a = table.cell(0, 0)#第一行第二个单元格
a.text = '性能指标'
b = table.cell(0, 1)#第一行第二个单元格
b.text = '响应速度'
document.save("第一章.docx")